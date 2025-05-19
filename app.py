from flask import Flask, render_template, jsonify, request, send_file
import asyncio
from generate_ai import ai_generate
import json
from database.db import Database
import io
from docx import Document
from docx.shared import Pt
import logging
import re

app = Flask(__name__)
db = Database()
logging.basicConfig(level=logging.INFO)

def clean_ai_response(response, response_type="lecture"):
    """Очищает ответ ИИ от markdown-обёртки и форматирования и пытается привести к валидному JSON
    
    Args:
        response (str): Ответ от ИИ
        response_type (str): Тип ответа ("lecture" или "programs")
    """
    try:
        clean = response.strip()
        
        # Проверяем, является ли ответ markdown-документом
        if clean.startswith('#') or '###' in clean:
            raise ValueError("Получен markdown-документ вместо JSON. Пожалуйста, проверьте формат ответа от ИИ.")
            
        # 1. Вырезаем только JSON-блок, если он есть
        match = re.search(r'```json\s*([\s\S]+?)\s*```', clean)
        if not match:
            match = re.search(r'```\s*([\s\S]+?)\s*```', clean)
        if match:
            clean = match.group(1).strip()
        else:
            # Если нет markdown-блока, ищем первую и последнюю фигурную скобку
            first = clean.find('{')
            last = clean.rfind('}')
            if first != -1 and last != -1 and last > first:
                clean = clean[first:last+1]
            else:
                raise ValueError("Не найден JSON-блок в ответе ИИ")
                
        # Логируем очищенный ответ для отладки
        logging.debug(f"Очищенный ответ: {clean}")
        
        # Попытка привести к валидному JSON
        # 1. Если ключи или строки в одинарных кавычках — заменить на двойные (если нет вложенных двойных)
        if "'" in clean and '"' not in clean:
            clean = clean.replace("'", '"')
        # 2. Если ключи без кавычек (YAML-стиль), добавить кавычки (простая эвристика)
        clean = re.sub(r'([,{]\s*)([a-zA-Z0-9_]+)(\s*:)','\\1"\\2"\\3', clean)
        # 3. Удалить лишние переносы строк внутри JSON
        clean = re.sub(r'\n+', ' ', clean)
        # 4. Удалить markdown-обёртки, если остались
        clean = clean.replace('```', '').strip()
        
        # Проверяем на незавершенные строки и исправляем их
        lines = clean.split('\n')
        fixed_lines = []
        for line in lines:
            if line.count('"') % 2 != 0:
                line = line + '"'
            fixed_lines.append(line)
        clean = '\n'.join(fixed_lines)
        
        # Проверяем и исправляем структуру JSON
        try:
            if not clean.startswith('{'):
                clean = '{' + clean
            if not clean.endswith('}'):
                clean = clean + '}'
            clean = re.sub(r'"\s*}\s*"', '", "', clean)
            clean = re.sub(r'"\s*}\s*{', '", "', clean)
            clean = re.sub(r',\s*,', ',', clean)
            clean = re.sub(r',\s*}', '}', clean)
            clean = re.sub(r',\s*]', ']', clean)
            
            # Пытаемся распарсить JSON
            result = json.loads(clean)
            
            # В зависимости от типа ответа применяем разные правила валидации
            if response_type == "lecture":
                # Проверяем наличие всех обязательных полей для лекции
                required_fields = ['introduction', 'sections', 'conclusion', 'recommendations']
                missing_fields = [field for field in required_fields if field not in result]
                
                if missing_fields:
                    # Если отсутствуют поля, добавляем их с пустыми значениями
                    for field in missing_fields:
                        if field == 'sections':
                            result[field] = []
                        elif field == 'recommendations':
                            result[field] = []
                        else:
                            result[field] = ""
            elif response_type == "programs":
                # Удаляем поля, специфичные для лекций
                lecture_fields = ['introduction', 'sections', 'conclusion', 'recommendations']
                for field in lecture_fields:
                    if field in result:
                        del result[field]
            
            return result
            
        except json.JSONDecodeError as e:
            logging.error(f"Ошибка при разборе JSON: {str(e)}")
            logging.error(f"Позиция ошибки: строка {e.lineno}, колонка {e.colno}")
            logging.error(f"Очищенный ответ: {clean}")
            raise ValueError(f"Не удалось разобрать JSON-ответ от ИИ: {str(e)}")
            
    except Exception as e:
        logging.error(f"Ошибка при очистке ответа: {str(e)}")
        logging.error(f"Исходный ответ: {response}")
        raise ValueError(f"Не удалось обработать ответ от ИИ: {str(e)}")

# --- Вспомогательная функция для безопасного вызова асинхронных функций ---
def safe_ai_generate_sync(prompt, mode, max_retries=3):
    async def inner():
        last_error = None
        for attempt in range(max_retries):
            try:
                result = await ai_generate(prompt, mode)
                if result is not None and result.strip():
                    return result
                print(f"⚠️ Пустой ответ от ИИ (попытка {attempt + 1}/{max_retries}), пробуем снова...")
            except Exception as e:
                last_error = e
                print(f"⚠️ Ошибка при генерации (попытка {attempt + 1}/{max_retries}): {str(e)}")
                if hasattr(e, 'response'):
                    print(f"Ответ API: {e.response.text}")
                    # Проверяем на превышение лимита
                    if 'Rate limit exceeded' in str(e.response.text):
                        raise ValueError("Превышен дневной лимит запросов к ИИ. Пожалуйста, попробуйте завтра или добавьте кредиты в настройках API.")
        if last_error:
            raise ValueError(f"❌ Не удалось получить корректный ответ от ИИ после {max_retries} попыток. Последняя ошибка: {str(last_error)}")
        raise ValueError(f"❌ Не удалось получить корректный ответ от ИИ после {max_retries} попыток.")
    return asyncio.run(inner())

@app.route('/')
def index():
    programs = db.get_all_programs()
    return render_template('index.html', programs=programs)

@app.route('/generate_programs', methods=['POST'])
def generate_programs():
    data = request.get_json()
    course_theme = data.get('course_theme', '')
    keywords = data.get('keywords', [])
    
    if not course_theme or not keywords:
        return jsonify({'error': 'Необходимо указать тему курса и ключевые слова'}), 400
    
    result = [course_theme, keywords]
    try:
        logging.info(f"Генерация программ для темы: {course_theme}, ключевые слова: {keywords}")
        programs = safe_ai_generate_sync(result, "names_programs")
        if not programs:
            raise ValueError("Получен пустой ответ от ИИ")
        programs_dict = clean_ai_response(programs, response_type="programs")
        logging.info(f"Сгенерированные программы: {programs_dict}")
        
        # Сохраняем программы в базу данных
        for title, description in programs_dict.items():
            db.save_program(title, description)
        
        return jsonify(programs_dict)
    except Exception as e:
        logging.exception('Ошибка при генерации программ:')
        return jsonify({'error': str(e)}), 500

@app.route('/generate_course_plan/<int:program_id>', methods=['POST'])
def generate_course_plan(program_id):
    program = db.get_program_by_id(program_id)
    if not program:
        logging.error(f'Программа с id={program_id} не найдена')
        return jsonify({'error': 'Программа не найдена'}), 404
    
    try:
        # Передаем и заголовок, и описание программы
        result = [program['title'], program['description']]
        plan = safe_ai_generate_sync(result, "generate_full_program")
        plan_dict = clean_ai_response(plan)
        logging.info(f'План курса для программы {program_id}: {plan_dict}')
        
        # Удаляем поля, специфичные для лекций, если они вдруг попали в план
        lecture_fields = ['introduction', 'sections', 'conclusion', 'recommendations']
        for field in lecture_fields:
            if field in plan_dict:
                del plan_dict[field]
        
        # Сортируем темы по их номерам
        sorted_plan = {}
        
        # Функция для извлечения номера темы
        def get_theme_number(theme):
            try:
                # Ищем число после слова "Тема" или "Тема:"
                match = re.search(r'Тема\s*:?\s*(\d+)', theme)
                if match:
                    return int(match.group(1))
                # Если это литература, помещаем в конец
                if theme.lower() == 'literature':
                    return float('inf')
                # Если не нашли номер, помещаем в конец
                return float('inf')
            except:
                return float('inf')
        
        # Сортируем темы по номеру
        sorted_themes = sorted(plan_dict.keys(), key=get_theme_number)
        
        # Добавляем темы в отсортированном порядке
        for theme in sorted_themes:
            if theme.lower() != 'literature':
                sorted_plan[theme] = plan_dict[theme]
        
        # В конце добавляем литературу, если она есть
        if 'literature' in plan_dict:
            sorted_plan['literature'] = plan_dict['literature']
        
        # Сохраняем план в базу данных
        db.save_course_plan(program_id, sorted_plan)
        return jsonify(sorted_plan)
    except Exception as e:
        logging.exception('Ошибка при генерации плана курса:')
        return jsonify({'error': str(e)}), 500

@app.route('/get_course_plan/<int:program_id>')
def get_course_plan(program_id):
    plan = db.get_course_plan(program_id)
    return jsonify(plan) if plan else jsonify({'error': 'План не найден'}), 404

@app.route('/update_course_plan/<int:program_id>', methods=['POST'])
def update_course_plan(program_id):
    data = request.get_json()
    db.update_course_plan(program_id, data)
    return jsonify({'success': True})

@app.route('/generate_lecture/<int:program_id>/<theme>', methods=['POST'])
def generate_lecture(program_id, theme):
    course_plan = db.get_course_plan(program_id)
    if not course_plan:
        logging.error(f'План курса для программы {program_id} не найден')
        return jsonify({'error': 'План курса не найден'}), 404
    
    program = db.get_program_by_id(program_id)
    if not program:
        logging.error(f'Программа с id={program_id} не найдена')
        return jsonify({'error': 'Программа не найдена'}), 404
    
    try:
        # Фильтруем литературу, если она есть
        if theme.lower() == 'literature':
            return jsonify({'error': 'Нельзя сгенерировать лекцию по литературе'}), 400
        if theme not in course_plan:
            return jsonify({'error': f'Тема {theme} не найдена в плане курса'}), 404
        
        # Подготавливаем данные для генерации
        theme_content = course_plan[theme]
        if not isinstance(theme_content, dict):
            logging.error(f'Неверный формат данных темы: {theme_content}')
            return jsonify({'error': 'Неверный формат данных темы'}), 500
            
        # Формируем результат в правильном формате для AI
        result = [
            program['title'],  # Название курса
            theme,            # Тема лекции
            course_plan,      # Структурированный план ВСЕЙ лекции
            theme_content     # Структурированный план необходимой лекции (пары)
        ]
        
        logging.info(f'Генерация лекции для темы {theme} с данными: {result}')
        
        # Добавляем явное указание формата в промпт
        prompt = f"""Сгенерируй лекцию по теме "{theme}" для курса "{program['title']}".
        Ответ должен быть в формате JSON со следующей структурой:
        {{
            "introduction": "Введение в тему",
            "sections": [
                {{
                    "title": "Название раздела",
                    "content": "Содержание раздела"
                }}
            ],
            "conclusion": "Заключение",
            "recommendations": ["Рекомендация 1", "Рекомендация 2"]
        }}
        
        Используй следующие данные для генерации:
        План курса: {course_plan}
        Содержание темы: {theme_content}
        """
        
        lecture = safe_ai_generate_sync(prompt, "generate_theme_lection")
        
        try:
            lecture_dict = clean_ai_response(lecture)
            # Если AI вернул словарь с одним ключом типа pair_1, pair_2 и т.д. — берём его содержимое
            if (
                isinstance(lecture_dict, dict)
                and len(lecture_dict) == 1
                and list(lecture_dict.keys())[0].startswith('pair_')
            ):
                lecture_dict = list(lecture_dict.values())[0]
            # Оставляем только нужные поля
            lecture_dict = {
                'introduction': lecture_dict.get('introduction', ''),
                'sections': lecture_dict.get('sections', []),
                'conclusion': lecture_dict.get('conclusion', ''),
                'recommendations': lecture_dict.get('recommendations', []),
            }
        except ValueError as e:
            logging.error(f"Ошибка при обработке ответа AI: {str(e)}")
            logging.error(f"Исходный ответ AI: {lecture}")
            return jsonify({'error': 'AI вернул ответ в неверном формате. Попробуйте сгенерировать лекцию снова.'}), 500
        
        # Проверяем структуру ответа
        if not isinstance(lecture_dict, dict):
            raise ValueError("Неверный формат ответа от ИИ")
        
        # Проверяем наличие обязательных полей
        required_fields = ['introduction', 'sections', 'conclusion', 'recommendations']
        missing_fields = [field for field in required_fields if field not in lecture_dict]
        if missing_fields:
            raise ValueError(f"В ответе отсутствуют обязательные поля: {', '.join(missing_fields)}")
        
        # Оборачиваем результат в ключ темы
        lecture_wrapped = {theme: lecture_dict}
        # Сохраняем лекцию в базу данных
        db.save_lecture(program_id, theme, lecture_wrapped)
        return jsonify(lecture_wrapped)
    except Exception as e:
        logging.exception('Ошибка при генерации лекции:')
        return jsonify({'error': str(e)}), 500

@app.route('/get_lecture/<int:program_id>/<theme>')
def get_lecture(program_id, theme):
    lecture = db.get_lecture(program_id, theme)
    return jsonify(lecture) if lecture else jsonify({'error': 'Лекция не найдена'}), 404

@app.route('/export_lecture/<int:program_id>/<theme>')
def export_lecture(program_id, theme):
    lecture = db.get_lecture(program_id, theme)
    if not lecture:
        return jsonify({'error': 'Лекция не найдена'}), 404
    
    # Получаем содержимое лекции
    # Если lecture[theme] содержит только нужные поля, используем их напрямую
    content = lecture.get(theme, lecture)
    
    # Создаем документ Word
    doc = Document()
    doc.add_heading(theme, 0)
    
    # Введение
    doc.add_heading('Введение', level=2)
    doc.add_paragraph(content.get('introduction', ''))
    
    # Основные разделы
    doc.add_heading('Основные разделы', level=2)
    sections = content.get('sections', [])
    if sections:
        for section in sections:
            doc.add_paragraph(section.get('title', ''), style='Heading 3')
            doc.add_paragraph(section.get('content', ''))
    else:
        doc.add_paragraph('Нет разделов для отображения', style='Intense Quote')
    
    # Заключение
    doc.add_heading('Заключение', level=2)
    doc.add_paragraph(content.get('conclusion', ''))
    
    # Рекомендации
    recommendations = content.get('recommendations', [])
    if recommendations:
        doc.add_heading('Рекомендации', level=2)
        for rec in recommendations:
            doc.add_paragraph(rec, style='List Bullet')
    
    # Сохраняем документ в память
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return send_file(
        doc_io,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'{theme}.docx'
    )

@app.route('/api/all_programs')
def api_all_programs():
    programs = db.get_all_programs()
    return jsonify(programs)

@app.route('/generate_big_lecture/<int:program_id>/<theme>', methods=['POST'])
def generate_big_lecture(program_id, theme):
    program = db.get_program_by_id(program_id)
    if not program:
        logging.error(f'Программа с id={program_id} не найдена')
        return jsonify({'error': 'Программа не найдена'}), 404
    try:
        # Формируем запрос для AI
        prompt = f"{theme} (курс: {program['title']})"
        logging.info(f'Генерация большой лекции по теме: {prompt}')
        lecture = safe_ai_generate_sync(prompt, "generate_big_lecture")
        lecture_dict = clean_ai_response(lecture)
        # Сохраняем лекцию в базу данных (можно в отдельную таблицу или как обычную лекцию)
        db.save_lecture(program_id, theme, lecture_dict)
        return jsonify(lecture_dict)
    except Exception as e:
        logging.exception('Ошибка при генерации большой лекции:')
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True) 